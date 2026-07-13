from pathlib import Path
from datetime import date
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "NeuroVibe_System_Architecture_and_Implementation_Plan.docx"
DIAGRAM = ROOT / "docs" / "neurovibe-architecture-diagram.png"

NAVY = "12324B"
BLUE = "2E74B5"
TEAL = "146C94"
GREEN = "27855B"
ORANGE = "CA6B21"
LIGHT_BLUE = "E8F2F8"
LIGHT_GREEN = "EDF8F1"
LIGHT_ORANGE = "FFF6EB"
LIGHT_GRAY = "F2F4F7"
MUTED = "557082"
CONTENT_WIDTH = 9360


def rgb(value):
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_font(run, size=11, color=NAVY, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("NeuroVibe System Architecture | Draft for Review | Page ")
    set_font(r, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    p._p.append(field)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=27, color=NAVY, bold=True)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
        r = p.add_run(subtitle)
        set_font(r, size=14, color=MUTED)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=11, color=TEAL, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_font(r, size=10.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_font(r, size=10, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=9.5, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def create_diagram():
    w, h = 1800, 1050
    im = Image.new("RGB", (w, h), "#F7FBFD")
    d = ImageDraw.Draw(im)
    bold = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 36)
    med = ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 23)
    body = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 20)
    small = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 17)
    d.rounded_rectangle((40, 30, 1760, 140), 22, fill="#12324B")
    d.text((80, 60), "NeuroVibe Connected Therapy System", font=bold, fill="white")
    d.text((80, 105), "Doctor-controlled care, simple patient use, dependable online and offline records", font=body, fill="#DCEEF5")

    boxes = [
        (70, 230, 480, 445, "Doctor / Admin", "Web dashboard\nRegister patients • assign devices\nSchedule sessions • review reports", "#E7F2FA", "#1E6F9F"),
        (70, 610, 480, 850, "Patient", "NeuroVibe mobile app\nPair device • follow schedule\nControl approved session • sync data", "#EDF8F1", "#27855B"),
        (660, 370, 1135, 675, "NeuroSense", "ESP32-C3 device\nBLE setup and app control\nWi-Fi server sync • 0–230 Hz\nOffline encrypted session queue", "#FFF6EB", "#CA6B21"),
        (1320, 235, 1730, 445, "Secure Backend API", "Authentication • assignment\nSchedules • session ingestion\nAudit and duplicate protection", "#F0EDFB", "#6B5BB4"),
        (1320, 610, 1730, 850, "Central Database", "Patients • devices • schedules\nSessions • status • reports\nDoctor-accessible clinical record", "#F1F5F7", "#597785"),
    ]
    for x1, y1, x2, y2, title, detail, fill, accent in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), 20, fill=fill, outline=accent, width=3)
        d.rounded_rectangle((x1+24, y1+25, x2-24, y1+78), 12, fill=accent)
        d.text((x1+44, y1+37), title, font=med, fill="white")
        y = y1 + 105
        for line in detail.split("\n"):
            d.text((x1+34, y), line, font=body if y == y1+105 else small, fill="#243F50")
            y += 31

    def arrow(points, color, label, label_pos):
        d.line(points, fill=color, width=7, joint="curve")
        x1, y1 = points[-2]
        x2, y2 = points[-1]
        import math
        a = math.atan2(y2-y1, x2-x1)
        s = 17
        left = (x2-s*math.cos(a-0.52), y2-s*math.sin(a-0.52))
        right = (x2-s*math.cos(a+0.52), y2-s*math.sin(a+0.52))
        d.polygon([(x2, y2), left, right], fill=color)
        d.text(label_pos, label, font=small, fill=color)

    arrow([(480, 335), (660, 430)], "#1E6F9F", "assignments, schedules, reports", (495, 300))
    arrow([(480, 730), (660, 610)], "#27855B", "BLE: setup, control, offline sync", (475, 692))
    arrow([(1135, 465), (1320, 345)], "#1E6F9F", "Wi-Fi HTTPS session upload", (1120, 390))
    arrow([(1525, 445), (1525, 610)], "#597785", "persisted records", (1545, 520))
    arrow([(480, 800), (880, 960), (1320, 750)], "#27855B", "App uploads records when phone internet returns", (660, 930))
    d.rounded_rectangle((70, 905, 1135, 1005), 16, fill="white", outline="#B6D0DC", width=2)
    d.text((98, 928), "Offline-first rule", font=med, fill="#146C94")
    d.text((98, 965), "The device keeps each completed session; Wi-Fi uploads directly, or BLE sends it to the app for later cloud synchronization.", font=small, fill="#355367")
    im.save(DIAGRAM)


def build_document():
    create_diagram()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    configure_styles(doc)
    add_footer(section)

    add_title(doc, "NeuroVibe System Architecture", "Implementation plan for NeuroSense ESP32-C3 device, patient app, doctor dashboard, and offline-first data synchronization")
    meta = [
        ("Document status", "Draft for architecture and senior review"),
        ("Version", "1.0"),
        ("Prepared for", "NeuroVibe project team"),
        ("Date", date.today().strftime("%d %B %Y")),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    set_table_geometry(table, [2450, 6910])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_font(r, size=10, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_font(r, size=10)
    doc.add_paragraph()
    add_callout(doc, "Purpose", "This document defines the recommended small-scale architecture for NeuroVibe. It prioritizes easy doctor and patient actions while preserving reliable records when the device, phone, or internet connection is unavailable.")
    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph("NeuroVibe is a connected vibration-therapy platform. A doctor assigns a NeuroSense device to a patient, sets a therapy schedule and permitted operating limits, and later reviews the usage data. The patient uses the NeuroVibe mobile app to connect to the device, follow a scheduled session, and control the device only within approved limits.")
    doc.add_paragraph("The platform is designed to be offline-first. NeuroSense records every completed session in its own memory. If Wi-Fi is available, it sends the record directly to the server. If Wi-Fi is unavailable but the patient app is connected by Bluetooth, it transfers the record to the app. The app stores the record securely and uploads it automatically when phone internet becomes available.")
    add_callout(doc, "Core principle", "The patient should only need to connect, follow the planned session, and press Start. The doctor handles patient registration, device assignment, schedules, replacement, exit decisions, and reporting. The system handles data synchronization automatically.")

    doc.add_heading("2. Scope and Assumptions", level=1)
    add_bullets(doc, [
        "Hardware currently in scope: ESP32-C3 Super Mini, DRV8833 motor driver, and two coin vibration motors.",
        "The device controls a requested vibration range from 0 Hz to 230 Hz. The PWM carrier remains separate from the requested vibration frequency.",
        "The current ERM coin motors provide estimated frequency control through calibrated PWM. A future accelerometer is required for measured, closed-loop frequency control.",
        "The patient receives a mobile app; the doctor/admin uses a browser dashboard built with HTML, CSS, and JavaScript.",
        "The first production-scale backend may be small, but it must support secure identity, patient-device assignment, session records, schedules, and synchronization acknowledgements.",
    ])

    doc.add_heading("3. Architecture Overview", level=1)
    doc.add_paragraph("The diagram below shows the permanent system record, the two communication routes, and the offline fallback path.")
    doc.add_picture(str(DIAGRAM), width=Inches(6.65))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("Figure 1. NeuroVibe architecture and synchronization paths")
    set_font(r, size=9, color=MUTED, italic=True)

    doc.add_heading("4. Entities, Roles, and Permissions", level=1)
    add_matrix(doc,
        ["Entity", "Primary platform", "Permitted actions", "Important restriction"],
        [
            ("Doctor / Admin", "Web dashboard", "Register patients; assign, replace, suspend, or retire devices; schedule sessions; review reports and status.", "Can only access patients and devices belonging to the doctor/organization."),
            ("Patient / User", "NeuroVibe mobile app", "Pair assigned device; receive reminders; run approved sessions; review own status; report a fault or exit request.", "Cannot exceed doctor-defined frequency or duration limits."),
            ("NeuroSense device", "ESP32-C3", "Execute approved motor commands; store session events; sync by Wi-Fi or BLE.", "Must accept commands only from the assigned, authenticated app/account."),
        ], [1500, 1850, 3600, 2410])

    doc.add_heading("5. NeuroSense Device Architecture", level=1)
    doc.add_heading("5.1 First-boot setup mode", level=2)
    add_steps(doc, [
        "At first boot, the device creates a unique device identifier and advertises a temporary Bluetooth name such as NeuroSense-8F4C91.",
        "The patient app, using an invite created by the doctor, authenticates and connects over BLE.",
        "The app sends the selected Wi-Fi network name and password, a device enrollment token, the assigned patient identifier, and initial therapy limits.",
        "The device stores configuration in encrypted non-volatile memory and confirms successful setup.",
        "If Wi-Fi connection succeeds, the device registers with the backend and downloads the latest permitted settings and schedule cache.",
    ])
    doc.add_heading("5.2 Firmware modules", level=2)
    add_matrix(doc, ["Module", "Responsibility"], [
        ("Motor controller", "Accepts only 0-230 Hz requests, converts a permitted target into PWM, controls both motors, and immediately stops on command or safety fault."),
        ("BLE service", "Supports setup, app authentication, control commands, schedule transfer, and offline session transfer."),
        ("Wi-Fi sync manager", "Uploads queued events by HTTPS, retries safely, and waits for server acknowledgement before marking an event as synchronized."),
        ("Local event queue", "Stores completed sessions in flash memory with unique session IDs while no upload route is available."),
        ("Configuration store", "Stores device identity, Wi-Fi settings, assignment metadata, permitted therapy profile, and cached schedule."),
        ("Safety manager", "Applies maximum frequency, duration, session state, watchdog, and emergency stop rules."),
    ], [2550, 6810])

    doc.add_heading("6. Patient Mobile App Architecture", level=1)
    doc.add_paragraph("The NeuroVibe app should be developed as an offline-first mobile application. Flutter is a practical choice because it can later support Android and iPhone from one codebase. The first release can be Android-focused if required.")
    add_matrix(doc, ["Screen", "Patient action", "System behavior"], [
        ("Device connection", "Tap Connect NeuroSense.", "Scans BLE, confirms device assignment, and displays connection status."),
        ("Today’s session", "Read the prescribed session and tap Start.", "Shows permitted frequency/duration and blocks out-of-plan commands."),
        ("Frequency control", "Use 0-230 Hz slider only where doctor permits manual control.", "Sends target frequency through BLE and displays requested Hz and percentage."),
        ("History and sync", "No technical action normally required.", "Shows local and synchronized session records; retries uploads automatically."),
        ("Support", "Report faulty device or request exit.", "Creates a request for the doctor dashboard without deleting historical data."),
    ], [1900, 2920, 4540])

    doc.add_heading("7. Doctor/Admin Web Dashboard", level=1)
    doc.add_paragraph("The dashboard must optimize for quick clinical operations. It should use a simple browser UI with clear patient search, assignment status, schedules, and reports rather than exposing device-level technical settings by default.")
    add_bullets(doc, [
        "Patient management: create, edit, suspend, and close patient records.",
        "Device inventory: available, assigned, faulty, lost, under repair, retired, and pending setup.",
        "Assignments: allocate a device, generate a patient invite/QR code, see assignment history, and initiate replacement.",
        "Scheduling: build a schedule template once and apply it to one or multiple patients.",
        "Monitoring: view latest synchronized session, last-seen time, pending sync count, usage/adherence, and device fault reports.",
        "Reporting: filter by patient, device, date, frequency, duration, and schedule adherence; export authorized reports.",
    ])

    doc.add_heading("8. Data Model", level=1)
    add_matrix(doc, ["Table", "Minimum purpose and key fields"], [
        ("doctors", "doctor_id, name, email, role, organization_id"),
        ("patients", "patient_id, doctor_id, patient_code, name, date_of_birth, gender, status"),
        ("devices", "device_id, display_name, status, firmware_version, last_seen_at"),
        ("device_assignments", "assignment_id, device_id, patient_id, assigned_by, starts_at, ends_at, status"),
        ("session_schedules", "schedule_id, patient_id, target_hz, maximum_hz, duration_seconds, recurrence, active"),
        ("therapy_sessions", "session_id, patient_id, device_id, schedule_id, started_at_utc, ended_at_utc, requested_hz, estimated_hz, pwm_value, sync_source"),
        ("device_status", "device_id, reported_at, connection_state, battery_level if available, pending_event_count"),
        ("audit_logs", "actor_id, action, entity_type, entity_id, timestamp, old_value, new_value"),
    ], [2500, 6860])
    add_callout(doc, "Data minimization", "Store a patient ID on NeuroSense whenever possible. Keep the patient name, demographic data, clinical history, and reports in the central database and encrypted mobile storage rather than permanently on the device.")

    doc.add_heading("9. Reliable Data Synchronization", level=1)
    doc.add_paragraph("Every completed session must receive a globally unique session ID on the device. The same ID follows the record through all routes. The server accepts the session once and ignores later duplicate uploads.")
    add_matrix(doc, ["Connection state", "First storage location", "Sync behavior"], [
        ("Device Wi-Fi available", "NeuroSense local queue", "Device uploads to API by HTTPS; record remains until the server acknowledges session_id."),
        ("No Wi-Fi, app connected by BLE", "NeuroSense then app encrypted database", "Device transfers event to app; app uploads immediately if it has internet, otherwise queues it."),
        ("No Wi-Fi and no phone internet", "NeuroSense and/or app local queue", "No data is deleted; automatic retry occurs at the next valid connection."),
        ("Both device and app upload", "Server", "Server uses session_id idempotency to keep one final clinical record."),
    ], [2600, 3000, 3760])

    doc.add_heading("10. Operational Scenarios", level=1)
    doc.add_heading("10.1 Scenario 1 - New patient receives a device", level=2)
    add_steps(doc, [
        "Doctor creates the patient record in the dashboard.",
        "Doctor selects an Available NeuroSense device and clicks Assign Device.",
        "System creates an assignment, a device setup invitation, and a QR/invite code.",
        "Patient installs NeuroVibe, signs in, scans the QR/invite code, and turns on the device.",
        "App pairs by BLE, sends Wi-Fi credentials and approved therapy settings, and confirms the device is active.",
        "Doctor creates or assigns a session schedule. The patient receives reminders in the app.",
    ])
    doc.add_heading("10.2 Scenario 2 - Existing patient needs a replacement device", level=2)
    add_steps(doc, [
        "Patient reports a fault through the app or contacts the doctor.",
        "Doctor marks the old device Faulty, Lost, Under Repair, or Retired and closes its assignment with an end date.",
        "Doctor assigns a replacement available device to the same patient.",
        "The system creates a new assignment and invite; previous records remain linked to the old device and patient.",
        "Patient pairs the replacement device. The app receives the existing schedule and permitted settings.",
    ])
    doc.add_heading("10.3 Scenario 3 - Patient exits the program", level=2)
    add_steps(doc, [
        "Patient selects Request Exit in the app. The app stops an active session before submitting the request.",
        "Doctor reviews the request, confirms closure, and records the closure reason if needed.",
        "System ends the active device assignment and disables new sessions for the patient.",
        "When the device is returned, it is reset, patient-specific configuration is securely erased, and its status changes to Available or Under Review.",
        "Historical session data remains in the central record according to the project’s approved data-retention policy.",
    ])

    doc.add_heading("11. Scheduling and Notifications", level=1)
    doc.add_paragraph("A doctor creates a schedule template once and assigns it to one or many patients. The backend creates a patient-specific scheduled-session record and sends a push notification to the patient app before the planned time.")
    add_matrix(doc, ["Schedule field", "Example"], [
        ("Template name", "Morning stimulation"),
        ("Permitted frequency", "150 Hz target; 0-180 Hz manual limit"),
        ("Duration", "10 minutes"),
        ("Recurrence", "Monday, Wednesday, Friday at 09:00"),
        ("Notification", "Reminder 15 minutes before the session"),
        ("Patients", "One patient or a selected group of patients"),
    ], [3000, 6360])
    add_callout(doc, "Offline schedule behavior", "The patient app caches the latest schedule. When the app connects to NeuroSense over BLE, it transfers the applicable schedule and limits to the device. The patient can still follow an already-downloaded schedule when internet is temporarily unavailable.")

    doc.add_heading("12. Security and Safety Baseline", level=1)
    add_bullets(doc, [
        "Use individual doctor and patient accounts; do not use one shared production password such as Doctor@1234.",
        "Require authentication before a patient can pair or control a device; confirm the patient-device assignment on every connection.",
        "Use TLS/HTTPS for server uploads and authenticated BLE pairing for setup and control.",
        "Encrypt Wi-Fi credentials and configuration on the device; encrypt patient records stored locally in the app.",
        "Never place database service keys or administrator secrets inside the mobile app or ESP32 firmware.",
        "Record audit events for patient creation, assignment, replacement, schedule changes, session deletion requests, and program closure.",
        "Apply a hard device safety maximum of 230 Hz and doctor-configurable session limits. Stop motors on disconnect, timeout, fault, or emergency stop.",
        "Treat the current frequency output as estimated until a vibration sensor/accelerometer validates actual motor vibration frequency.",
    ])

    doc.add_heading("13. Recommended Implementation Roadmap", level=1)
    add_matrix(doc, ["Phase", "Deliverable", "Success condition"], [
        ("1. Foundation", "Database, authentication, doctor/patient roles, device inventory.", "Doctor can create patient and see an available device."),
        ("2. Assignment", "Dashboard assignment flow, QR/invite generation, assignment history.", "A patient has one active assigned device."),
        ("3. Device setup", "ESP32 ID, BLE provisioning, encrypted configuration, local event queue.", "Patient can pair and provision Wi-Fi without developer tools."),
        ("4. Patient app", "Connection, 0-230 Hz control, start/stop, local encrypted session storage.", "Patient can complete an approved offline session."),
        ("5. Synchronization", "Wi-Fi direct upload, BLE transfer, phone upload, idempotent server acknowledgement.", "No session is lost or duplicated across reconnection."),
        ("6. Operations", "Schedules, notifications, replacement, exit workflow, reports, audit logs.", "Doctor can manage a multi-patient program from the web dashboard."),
        ("7. Measurement", "Accelerometer-based frequency verification and calibration.", "System distinguishes requested, estimated, and measured vibration frequency."),
    ], [1150, 4350, 3860])

    doc.add_heading("14. Acceptance Criteria for the First Working Prototype", level=1)
    add_bullets(doc, [
        "Doctor can add a patient, assign a unique NeuroSense device, and create a scheduled session.",
        "Patient can use NeuroVibe to pair with the assigned device and run a permitted session between 0 and 230 Hz.",
        "Device creates a timestamped session record with device ID, patient ID, requested frequency, duration, and session ID.",
        "Device uploads through Wi-Fi when available; otherwise it retains the session record.",
        "App receives a device record through BLE, stores it locally when offline, and uploads it when internet returns.",
        "Doctor dashboard shows synchronized patient history and the last known device status.",
        "Faulty-device replacement preserves historical records and creates a new assignment rather than overwriting the old one.",
        "Program exit ends the assignment and securely clears patient-specific configuration from a returned device.",
    ])

    doc.add_heading("15. Open Decisions for Senior Review", level=1)
    add_bullets(doc, [
        "Confirm whether NeuroVibe will launch Android-only or cross-platform with Flutter for Android and iPhone.",
        "Confirm database/backend provider and required hosting region.",
        "Define clinical schedule rules: maximum duration, allowable frequency ranges, missed-session behavior, and emergency-stop procedure.",
        "Confirm whether actual vibration frequency must be measured. If yes, select an accelerometer and add closed-loop control in a later phase.",
        "Define approved data-retention, consent, export, deletion, and privacy requirements before using real patient data.",
        "Decide whether direct remote control is permitted. The recommended first version limits live control to nearby BLE connection for safety.",
    ])

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run("End of architecture document")
    set_font(r, size=9, color=MUTED, italic=True)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
