from pathlib import Path
from datetime import date
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_neurovibe_architecture import (
    NAVY, BLUE, TEAL, GREEN, ORANGE, LIGHT_GRAY, MUTED,
    configure_styles, add_footer, add_title, add_callout, add_bullets,
    add_steps, add_matrix, set_font, set_table_geometry, set_cell_shading,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "NeuroVibe_Replanned_Architecture_v2.docx"
ARCH = ROOT / "docs" / "neurovibe-replanned-architecture-v2.png"
LIFE = ROOT / "docs" / "neurovibe-lifecycle-v2.png"


def fonts():
    return {
        "title": ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 38),
        "head": ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 25),
        "body": ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 19),
        "small": ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 16),
        "small_bold": ImageFont.truetype("C:/Windows/Fonts/arialbd.ttf", 16),
    }


def draw_arrow(draw, start, end, color, width=6):
    import math
    draw.line([start, end], fill=color, width=width)
    a = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 18
    p1 = (end[0] - size * math.cos(a - 0.48), end[1] - size * math.sin(a - 0.48))
    p2 = (end[0] - size * math.cos(a + 0.48), end[1] - size * math.sin(a + 0.48))
    draw.polygon([end, p1, p2], fill=color)


def draw_box(draw, rect, title, lines, fill, border, f):
    x1, y1, x2, y2 = rect
    draw.rounded_rectangle(rect, 18, fill=fill, outline=border, width=3)
    draw.rounded_rectangle((x1 + 18, y1 + 18, x2 - 18, y1 + 66), 10, fill=border)
    draw.text((x1 + 34, y1 + 28), title, font=f["head"], fill="white")
    y = y1 + 86
    for line in lines:
        draw.text((x1 + 28, y), line, font=f["small"], fill="#243F50")
        y += 26


def create_architecture_diagram():
    f = fonts()
    image = Image.new("RGB", (1900, 1250), "#F7FBFD")
    d = ImageDraw.Draw(image)
    d.rounded_rectangle((35, 25, 1865, 130), 20, fill="#12324B")
    d.text((75, 50), "NeuroVibe V2 — Replanned End-to-End Architecture", font=f["title"], fill="white")
    d.text((75, 98), "Server-owned truth • Simple user actions • Durable offline synchronization", font=f["body"], fill="#DCEEF5")

    draw_box(d, (60, 205, 500, 430), "Doctor Web Portal", [
        "Create patient and care plan",
        "Assign / replace / retire device",
        "Schedule groups and review data",
        "Approve exit and manage consent",
    ], "#E7F2FA", "#1E6F9F", f)
    draw_box(d, (60, 665, 500, 900), "Patient NeuroVibe App", [
        "Accept invitation and pair device",
        "Provision Wi-Fi through BLE",
        "Run one-tap approved sessions",
        "Cache and relay offline events",
    ], "#EDF8F1", "#27855B", f)
    draw_box(d, (690, 390, 1210, 715), "NeuroSense ESP32-C3", [
        "Unique immutable device ID",
        "BLE commissioning and control",
        "Wi-Fi HTTPS synchronization",
        "Local durable session queue",
        "0–230 Hz hard safety boundary",
        "Stores patient ID, not full profile",
    ], "#FFF6EB", "#CA6B21", f)
    draw_box(d, (1395, 195, 1830, 430), "Backend API", [
        "Authentication and authorization",
        "Assignment and schedule rules",
        "Idempotent session ingestion",
        "Audit trail and notifications",
    ], "#F0EDFB", "#6B5BB4", f)
    draw_box(d, (1395, 665, 1830, 900), "PostgreSQL", [
        "Patients and program status",
        "Devices and assignment history",
        "Schedules and therapy sessions",
        "Audit and synchronization receipts",
    ], "#F1F5F7", "#597785", f)

    draw_arrow(d, (500, 305), (1395, 305), "#1E6F9F")
    d.text((700, 270), "HTTPS: patients • assignments • schedules • reports", font=f["small_bold"], fill="#1E6F9F")
    draw_arrow(d, (1395, 370), (500, 370), "#1E6F9F")
    draw_arrow(d, (500, 785), (690, 620), "#27855B")
    d.text((505, 700), "BLE: setup • control • offline relay", font=f["small_bold"], fill="#27855B")
    draw_arrow(d, (690, 680), (500, 850), "#27855B")
    draw_arrow(d, (1210, 475), (1395, 390), "#CA6B21")
    d.text((1190, 430), "Wi-Fi HTTPS", font=f["small_bold"], fill="#CA6B21")
    draw_arrow(d, (1610, 430), (1610, 665), "#597785")
    d.text((1635, 530), "validated records", font=f["small_bold"], fill="#597785")
    # Route phone uploads to the API (never directly to PostgreSQL).
    d.line([(500, 875), (1270, 950), (1270, 470), (1395, 420)], fill="#27855B", width=6, joint="curve")
    draw_arrow(d, (1270, 470), (1395, 420), "#27855B")
    d.text((690, 915), "App uploads encrypted queue to API when internet returns", font=f["small_bold"], fill="#27855B")

    d.rounded_rectangle((60, 1015, 1830, 1175), 18, fill="white", outline="#B7CDD7", width=2)
    d.text((90, 1045), "One session, one event ID, two possible upload routes", font=f["head"], fill="#146C94")
    d.text((90, 1090), "NeuroSense always records first. It uploads directly by Wi-Fi or transfers the same event to NeuroVibe by BLE.", font=f["body"], fill="#355367")
    d.text((90, 1125), "The server uses the event ID to acknowledge one authoritative record and reject duplicates without losing data.", font=f["body"], fill="#355367")
    image.save(ARCH)


def create_lifecycle_diagram():
    f = fonts()
    image = Image.new("RGB", (1900, 1150), "#FAFCFD")
    d = ImageDraw.Draw(image)
    d.text((55, 45), "NeuroVibe Device and Patient Lifecycle", font=f["title"], fill="#12324B")
    d.text((55, 95), "Historical assignments are closed, never overwritten", font=f["body"], fill="#557082")

    def lane(y, title, labels, colors):
        d.text((55, y), title, font=f["head"], fill="#12324B")
        x = 230
        boxes = []
        for label, color in zip(labels, colors):
            rect = (x, y - 15, x + 220, y + 65)
            d.rounded_rectangle(rect, 15, fill="#FFFFFF", outline=color, width=4)
            bbox = d.textbbox((0, 0), label, font=f["small_bold"])
            tx = x + (220 - (bbox[2] - bbox[0])) / 2
            d.text((tx, y + 10), label, font=f["small_bold"], fill="#243F50")
            boxes.append(rect)
            x += 275
        for a, b in zip(boxes, boxes[1:]):
            draw_arrow(d, (a[2] + 10, y + 25), (b[0] - 10, y + 25), "#6B8795", 5)

    lane(205, "Device", ["Factory New", "Claimed", "Assigned", "Active", "Faulty / Returned", "Sanitized / Available"], ["#6B8795", "#1E6F9F", "#6B5BB4", "#27855B", "#C24B58", "#CA6B21"])
    lane(470, "Patient Program", ["Invited", "Enrolled", "Active", "Paused", "Exit Requested", "Closed"], ["#1E6F9F", "#6B5BB4", "#27855B", "#CA6B21", "#C24B58", "#597785"])
    lane(735, "Session", ["Scheduled", "Ready", "Running", "Completed", "Queued Offline", "Acknowledged"], ["#1E6F9F", "#6B5BB4", "#CA6B21", "#27855B", "#C78B2F", "#597785"])
    d.rounded_rectangle((55, 940, 1840, 1085), 18, fill="#F2F7FA", outline="#B7CDD7", width=2)
    d.text((85, 972), "Replacement rule", font=f["head"], fill="#146C94")
    d.text((85, 1015), "Close the old assignment → preserve its sessions → sanitize or retire the old device → create a new assignment.", font=f["body"], fill="#355367")
    d.text((85, 1052), "Patient exit closes future schedules and device access, but historical records follow the approved retention policy.", font=f["body"], fill="#355367")
    image.save(LIFE)


def add_picture_with_alt(doc, path, width, title, description):
    shape = doc.add_picture(str(path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)
    return shape


def build():
    create_architecture_diagram()
    create_lifecycle_diagram()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    configure_styles(doc)
    add_footer(section)

    add_title(doc, "NeuroVibe Replanned Architecture", "Version 2 — complete product, device, backend, data, and operational design")
    meta = [
        ("Status", "Replanned draft for senior review"),
        ("Design objective", "Simple doctor and patient actions; robust backend and offline behavior"),
        ("Platforms", "Doctor web portal, NeuroVibe patient app, NeuroSense ESP32-C3"),
        ("Prepared", date.today().strftime("%d %B %Y")),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    set_table_geometry(table, [2300, 7060])
    for row, (label, value) in zip(table.rows, meta):
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), size=10, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value), size=10)
    doc.add_paragraph()
    add_callout(doc, "Replanning decision", "The central server is the source of truth. NeuroSense and NeuroVibe are durable offline clients. Device assignment, patient enrollment, therapy schedules, session records, replacement, and exit are separate historical lifecycles rather than one editable device record.")
    doc.add_page_break()

    doc.add_heading("1. Product Definition", level=1)
    doc.add_paragraph("NeuroVibe is a doctor-managed, patient-operated vibration-session system. The doctor creates the care relationship and defines what the patient may do. The patient uses a dedicated app to commission the assigned NeuroSense device and run scheduled sessions. NeuroSense records device-generated usage events, even when both Wi-Fi and phone internet are unavailable.")
    add_callout(doc, "User-experience target", "Doctor: Create patient → Assign device → Schedule sessions → Review results. Patient: Accept invite → Connect device → Start scheduled session. All technical synchronization happens in the background.")

    doc.add_heading("2. Correct Domain Model", level=1)
    doc.add_paragraph("Doctor and patient are user roles, not the only database entities. A reliable system also requires devices, assignments, care plans, schedules, sessions, synchronization receipts, and audit records.")
    add_matrix(doc, ["Domain object", "Purpose", "Lifecycle owner"], [
        ("Doctor account", "Authenticates the clinician/admin and scopes accessible patients.", "Organization/admin"),
        ("Patient account/profile", "Represents the user, consent state, and program status.", "Doctor with patient participation"),
        ("Device", "Permanent hardware identity and current inventory condition.", "Organization"),
        ("Device assignment", "Time-bounded relationship between one patient and one device.", "Doctor"),
        ("Care plan", "Doctor-approved frequency range, duration, and control rules.", "Doctor"),
        ("Schedule", "Specific or recurring sessions and reminder timing.", "Doctor"),
        ("Therapy session", "Immutable usage record generated by the device/app workflow.", "System"),
        ("Audit event", "Who changed assignment, schedule, status, or access and when.", "System"),
    ], [2200, 4700, 2460])

    doc.add_heading("3. Replanned System Architecture", level=1)
    add_picture_with_alt(
        doc, ARCH, Inches(6.65),
        "NeuroVibe replanned architecture",
        "Doctor web portal and patient app communicate with a secure backend API. NeuroSense uses Wi-Fi for direct session upload and Bluetooth for app setup, control, and offline relay. The backend validates records before storing them in PostgreSQL.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Figure 1. Replanned online and offline architecture"), size=9, color=MUTED, italic=True)
    add_bullets(doc, [
        "Doctor web portal and patient app never connect directly to the database. Both use the authenticated backend API.",
        "NeuroSense can upload through Wi-Fi, while NeuroVibe provides a Bluetooth relay when direct Wi-Fi is unavailable.",
        "The same session event ID is preserved through both routes, so duplicate uploads are harmless.",
        "Only the server confirms that a record is accepted. Local copies are retained until acknowledgement.",
    ])

    doc.add_heading("4. Technology Boundaries for a Small-Scale Build", level=1)
    add_matrix(doc, ["Layer", "Recommended first implementation", "Reason"], [
        ("Doctor portal", "Responsive HTML/CSS/JavaScript frontend; React or vanilla JS for the first version.", "Runs in a browser and keeps doctor workflow simple."),
        ("Patient app", "Flutter Android-first application.", "Supports BLE, local encrypted storage, notifications, and later iOS support."),
        ("Backend", "Single modular API service; FastAPI or Node/NestJS.", "One deployable backend is sufficient; no microservices are needed initially."),
        ("Database", "Managed PostgreSQL with authentication and row-level authorization.", "Fits assignments, schedules, session history, and reports."),
        ("Notifications", "Mobile push notification service plus in-app schedule cache.", "Reminders can reach many patients without the ESP32 being online."),
        ("Device", "ESP32-C3 firmware with BLE, HTTPS, encrypted settings, and flash event queue.", "Meets commissioning, control, and offline requirements."),
    ], [1750, 4800, 2810])

    doc.add_heading("5. Identity, Enrollment, and Device Commissioning", level=1)
    doc.add_heading("5.1 Device identity", level=2)
    add_bullets(doc, [
        "Create an immutable device ID at manufacturing or first boot from a secure random UUID; never derive identity only from a Bluetooth name.",
        "Use a friendly name such as NeuroSense-A7F3 for discovery, while the server stores the full device UUID.",
        "A new device begins Factory New and cannot run patient sessions until it is claimed and assigned.",
        "Use a QR code or short claim code printed on the device/packaging to prevent assignment of the wrong nearby device.",
    ])
    doc.add_heading("5.2 Patient enrollment", level=2)
    add_steps(doc, [
        "Doctor creates a patient and selects Assign NeuroSense.",
        "Server creates a one-time invitation linked to the intended patient, device, and doctor.",
        "Patient signs in to NeuroVibe and accepts the invitation or scans the QR code.",
        "App verifies that the discovered BLE device ID matches the server assignment.",
        "Patient selects the home/clinic Wi-Fi network and enters the password in the app.",
        "App sends Wi-Fi credentials and a short-lived device enrollment credential over authenticated BLE.",
        "Device connects to the backend, exchanges the enrollment credential for its own device credential, and reports Active.",
    ])
    add_callout(doc, "Privacy boundary", "Do not send the patient’s full name, age, or gender to NeuroSense unless there is a demonstrated offline requirement. Store patient_id, assignment_id, and care-plan limits on the device; keep identifiable profile data on the server and encrypted app storage.")

    doc.add_heading("6. Therapy and Frequency-Control Model", level=1)
    add_matrix(doc, ["Value", "Meaning"], [
        ("requested_hz", "Frequency requested by the approved care plan or app control."),
        ("estimated_hz", "Frequency inferred from calibrated PWM when no vibration sensor exists."),
        ("measured_hz", "Actual vibration frequency measured by a future accelerometer; null in the initial prototype."),
        ("pwm_value", "0-255 motor command used by the current 8-bit ESP32 PWM configuration."),
        ("safety maximum", "Hard device and server ceiling of 230 Hz; doctor may prescribe a lower maximum."),
    ], [2300, 7060])
    doc.add_paragraph("The existing ESP32-C3 + DRV8833 + ERM coin-motor configuration cannot guarantee exact vibration frequency without feedback. PWM controls average motor voltage and speed; the software should describe output as estimated until calibration or accelerometer feedback is added.")
    add_bullets(doc, [
        "Keep the PWM carrier frequency separate from vibration frequency.",
        "Reject commands below 0 Hz or above 230 Hz rather than silently accepting them.",
        "Apply doctor-defined target, maximum frequency, maximum duration, cooldown, and allowed-time rules on both server/app and device.",
        "Stop immediately on timeout, BLE control loss where required, watchdog failure, or patient emergency stop.",
    ])

    doc.add_heading("7. Session Lifecycle", level=1)
    add_steps(doc, [
        "Backend creates a scheduled session from the doctor’s care plan and schedule.",
        "Patient app downloads and caches the schedule, then sends the applicable permitted session to NeuroSense over BLE.",
        "At start, NeuroSense creates session_id and records device_id, assignment_id, schedule_id, timestamps, and requested settings.",
        "During operation, the device applies safety limits and records relevant setting changes or summary values.",
        "At stop/completion, NeuroSense closes the session event and saves it to durable flash before reporting success to the app.",
        "The session uploads by device Wi-Fi or app relay. The backend validates assignment and schedule context, stores it once, and returns acknowledgement.",
        "Device/app mark the event acknowledged and later remove it according to local retention limits.",
    ])

    doc.add_heading("8. Offline-First Synchronization Contract", level=1)
    add_matrix(doc, ["Rule", "Required behavior"], [
        ("Record first", "Write a completed session to durable local storage before attempting network transmission."),
        ("Immutable identity", "Every event has session_id, device_id, assignment_id, sequence number, and created_at."),
        ("Idempotent upload", "POSTing the same session_id again returns the original acknowledgement instead of creating a duplicate."),
        ("Acknowledgement", "Delete/compact local data only after the server confirms the event ID."),
        ("Ordering", "Use device sequence numbers to detect gaps even if events arrive out of order."),
        ("Conflict rule", "Device-generated session facts are append-only; server annotations never rewrite raw device facts."),
        ("Time quality", "Store UTC timestamp plus source: NTP, phone-provided, or unsynchronized. Preserve sequence when time is uncertain."),
    ], [2150, 7210])

    doc.add_heading("9. Doctor and Patient Experience", level=1)
    doc.add_heading("9.1 Doctor portal", level=2)
    add_matrix(doc, ["Doctor goal", "Primary action", "Result"], [
        ("Enroll patient", "Add Patient", "Creates patient and sends invitation."),
        ("Allocate hardware", "Assign Device", "Creates a historical assignment and claim code."),
        ("Plan care", "Create Schedule", "Selects frequency/duration template and one or many patients."),
        ("Monitor", "Open Patient", "Shows adherence, sessions, last sync, alerts, and assigned device."),
        ("Replace hardware", "Replace Device", "Closes old assignment and begins replacement wizard."),
        ("Close participation", "Close Program", "Ends schedules and access after confirmation."),
    ], [2100, 2050, 5210])
    doc.add_heading("9.2 Patient app", level=2)
    add_bullets(doc, [
        "Home screen shows one clear card: next session time and Start Session button.",
        "Connection screen shows Connected, Setup required, or Device unavailable in plain language.",
        "Manual frequency slider appears only if the doctor’s care plan allows it and remains within the prescribed range.",
        "Sync status uses simple messages such as All data saved or Waiting for internet; no technical queue controls are exposed.",
        "Report Device Problem and Request Exit are visible support actions, not destructive local deletion actions.",
    ])

    doc.add_heading("10. Replanned Operational Scenarios", level=1)
    doc.add_heading("10.1 New patient and device assignment", level=2)
    add_steps(doc, [
        "Doctor adds the patient and care plan.",
        "Doctor selects an Available device and confirms Assign.",
        "Patient receives an app invitation and scans the device claim code.",
        "App pairs through BLE and provisions Wi-Fi.",
        "Device authenticates with the backend and downloads limits.",
        "Portal shows Assignment Active; patient sees the next scheduled session.",
    ])
    doc.add_heading("10.2 Faulty-device replacement", level=2)
    add_steps(doc, [
        "Patient taps Report Device Problem; the portal creates an actionable alert.",
        "Doctor chooses Replace Device and selects a replacement from Available inventory.",
        "Backend closes the old assignment at an exact timestamp without changing prior sessions.",
        "Old device status becomes Faulty/Return Pending and cannot upload new patient sessions after closure except queued events created before closure.",
        "A new assignment and claim invitation are created for the replacement device.",
        "Patient commissions the replacement; existing care plan and future schedule remain attached to the patient, not copied from the faulty hardware.",
        "Returned hardware is synchronized, sanitized, inspected, and then marked Available or Retired.",
    ])
    doc.add_heading("10.3 Patient exits the test/program", level=2)
    add_steps(doc, [
        "Patient submits Request Exit; the app confirms that this stops future sessions but does not erase records immediately.",
        "Doctor reviews and confirms Close Program.",
        "Backend cancels future schedules, revokes patient control authorization, and closes the active device assignment.",
        "Device receives a decommission command when next online or through the app, uploads remaining events, and clears patient-specific configuration.",
        "Historical records remain available only according to the approved consent and retention policy.",
        "Device returns to Sanitized/Available or Retired after inspection.",
    ])

    doc.add_heading("11. Group Scheduling and Notifications", level=1)
    add_steps(doc, [
        "Doctor creates a reusable schedule template with frequency, duration, recurrence, start/end dates, and reminder lead time.",
        "Doctor selects multiple eligible patients; the portal previews conflicts and each patient’s permitted limits.",
        "Backend creates patient-specific schedule instances rather than one shared mutable record.",
        "Notification service sends each patient a reminder; the app caches upcoming sessions.",
        "Completion, missed, cancelled, or pending-sync status is recorded per patient.",
        "Doctor sees group adherence while retaining patient-level drill-down and privacy boundaries.",
    ])

    doc.add_heading("12. Data Model and Key Constraints", level=1)
    add_matrix(doc, ["Table", "Key fields / constraint"], [
        ("users", "user_id, role, organization_id, status; unique login identity"),
        ("patients", "patient_id, doctor_id, profile, consent_status, program_status"),
        ("devices", "device_id, serial/display name, lifecycle_status, firmware, last_seen"),
        ("device_assignments", "assignment_id, device_id, patient_id, starts_at, ends_at; no overlapping active assignment"),
        ("care_plans", "care_plan_id, patient_id, min_hz, max_hz <= 230, max_duration, effective dates"),
        ("schedule_templates", "template parameters owned by doctor/organization"),
        ("scheduled_sessions", "patient-specific occurrence, planned time, status, reminder status"),
        ("therapy_sessions", "session_id unique, assignment_id, schedule_id, requested/estimated/measured Hz, duration, sync source"),
        ("sync_receipts", "event_id unique, received_at, route, device sequence, acknowledgement"),
        ("audit_events", "actor, action, target, before/after summary, UTC timestamp"),
    ], [2400, 6960])

    doc.add_heading("13. API Surface", level=1)
    add_matrix(doc, ["Area", "Example operations"], [
        ("Authentication", "login, refresh token, logout, password reset"),
        ("Patients", "create patient, update profile, invite, pause/close program"),
        ("Devices", "claim, list inventory, report status, mark faulty, sanitize, retire"),
        ("Assignments", "assign, replace, close, get current and history"),
        ("Care plans", "create version, activate, expire; device/app read effective plan"),
        ("Schedules", "create template, assign patients, list occurrences, cancel"),
        ("Sessions", "idempotent upload, acknowledgement, patient history, doctor report"),
        ("Sync", "upload batch, acknowledge IDs, report sequence gaps"),
    ], [2400, 6960])

    doc.add_heading("14. Security, Privacy, and Safety", level=1)
    add_bullets(doc, [
        "Use individual doctor and patient accounts with role- and organization-scoped authorization.",
        "Use a unique device credential per NeuroSense unit; never embed a shared database or administrator key in firmware.",
        "Authenticate BLE commissioning with a one-time claim code and short-lived enrollment credential.",
        "Encrypt transport, device configuration, and mobile local records. Keep Wi-Fi passwords out of server logs and therapy records.",
        "Keep patient-identifying data off the ESP32 unless explicitly required and approved.",
        "Use audit events for access, assignment, care-plan changes, replacement, exit, export, and deletion actions.",
        "Enforce frequency and duration limits independently on backend, app, and device; the device remains the final safety boundary.",
        "For clinical deployment, obtain formal privacy, cybersecurity, safety, and regulatory review before using real patient data.",
    ])

    doc.add_heading("15. Lifecycle Model", level=1)
    add_picture_with_alt(
        doc, LIFE, Inches(6.65),
        "NeuroVibe lifecycle model",
        "Three independent state flows show device lifecycle, patient-program lifecycle, and therapy-session lifecycle. Historical assignments are closed rather than overwritten.",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Figure 2. Independent device, patient-program, and session lifecycles"), size=9, color=MUTED, italic=True)

    doc.add_heading("16. Implementation Roadmap", level=1)
    add_matrix(doc, ["Phase", "Build", "Exit criteria"], [
        ("1. Backend core", "Auth, PostgreSQL schema, API, audit and idempotent session ingestion.", "Test users, assignments, schedules, and duplicate-safe uploads work."),
        ("2. Doctor portal", "Patient, inventory, assignment, care plan and scheduling screens.", "Doctor completes scenario 1 without developer assistance."),
        ("3. Device foundation", "Identity, BLE commissioning, encrypted settings, motor safety, flash event queue.", "Device survives reboot/offline use without losing a session."),
        ("4. Patient app", "Invite, BLE setup, session UI, encrypted local queue and notifications.", "Patient completes a scheduled offline session in a few taps."),
        ("5. End-to-end sync", "Device Wi-Fi upload, app relay, acknowledgement, sequence-gap handling.", "Reconnect tests produce no lost or duplicated sessions."),
        ("6. Operations", "Replacement, exit, group schedules, reports, sanitization flow.", "All three operational scenarios pass acceptance testing."),
        ("7. Measurement", "Accelerometer calibration and closed-loop control if required.", "Actual measured Hz is validated across the intended range."),
    ], [1150, 4620, 3590])

    doc.add_heading("17. Senior Review Decisions", level=1)
    add_bullets(doc, [
        "Approve patient app platform: Android-first Flutter or Android+iOS launch.",
        "Approve initial backend stack and hosting/data region.",
        "Approve whether NeuroSense may run only scheduled sessions or also doctor-authorized manual sessions.",
        "Approve clinical limits, emergency-stop behavior, and definition of a completed/missed session.",
        "Decide whether estimated motor frequency is sufficient for the prototype or an accelerometer is mandatory.",
        "Approve consent, retention, export, deletion, and device-sanitization policies before real patient enrollment.",
    ])

    doc.add_heading("18. Prototype Acceptance Checklist", level=1)
    add_bullets(doc, [
        "One doctor can enroll patients, assign unique devices, and schedule individual or group sessions.",
        "A patient can commission only the device assigned to their invitation.",
        "NeuroSense enforces 0-230 Hz and doctor-defined limits even without the app or internet.",
        "Every completed session is durable across reboot and loss of connectivity.",
        "Wi-Fi direct upload and BLE app relay create one server record, not duplicates.",
        "Doctor can replace a device without losing history or moving old sessions to the new device.",
        "Patient exit disables future operation, uploads remaining data, and clears patient-specific device configuration.",
        "Doctor can see last synchronized status and distinguish live, stale, pending, missed, and completed information.",
    ])

    DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
